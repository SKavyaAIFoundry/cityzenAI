import docx2txt                                         # adapted from docx
import os
import os.path
import pypandoc                                         # MIT License
import re

from docx import Document   # pip install python-docx   # MIT License
from msilib.schema import Directory

from utilities.general_utilities import resetDirectory


# methods for extracting content from Word files

def analyseDocxFileDocx(fileName):
    # https://stackoverflow.com/questions/41200127/extract-image-position-from-docx-file-using-python-docx

    word_count = 0
    paragraph_count = 0
    bullet_count = 0
    image_count = 0

    # extract text from all document paragraphs
    document = Document(fileName)
    document_text = '\n\n'.join(
        paragraph.text for paragraph in document.paragraphs
    )

    # split text into words to get word count
    document_words = document_text.split()
    word_count = len(document_words)

    for paragraph in document.paragraphs:
        if any (paragraph.text):
            paragraph_count += 1

    for image in document.inline_shapes:
        image_count += 1
    # docx does not currently support floating shapes/images :(

    return document_text, word_count, paragraph_count, bullet_count, image_count


def analyseDocxFileDocx2txt(fileName):
    # https://pypi.org/project/docx2txt/

    word_count = 0
    paragraph_count = 0
    bullet_count = 0
    image_count = 0

    document = Document(fileName)
    for paragraph in document.paragraphs:
        if any (paragraph.text):
            paragraph_count += 1

    # works ok but not 100% accurate as it is working with MS Word styles
    #for paragraph in document.paragraphs:
        if paragraph.style.name == 'List Paragraph':
            # add to count if the paragraph is not empty
            if len(paragraph.text) > 0:
                bullet_count += 1
                print(str(bullet_count) + " - " + str(paragraph.text))
    #print("bullets: " + str(bullet_count))

    #for paragraph in document.paragraphs:
        # look for numbered sections which are not bulleted e.g., '1.\t'
        matches = re.finditer(r'(?=(\d+[.]\t))', paragraph.text)
        results = [(match.group(1)) for match in matches]
        #print(len(results))
        #print(results)
        bullet_count += len(results)
    print("bullets: " + str(bullet_count))

    # use this to just extract the text
    #text = docx2txt.process(fileName)
    
    # use this to extract the text and image files (to a specified directory)

    # delete everything in the temporary image directory
    image_path = './temp_image_files'
    resetDirectory(image_path)
    # extract the text and images
    document_text = docx2txt.process(fileName, image_path)
    document_words = document_text.split()
    word_count = len(document_words)

    # count the number of unique images in the doc
    
    # simple version for working with CWD
    #image_count = (len([name for name in os.listdir('.') if os.path.isfile(name)]))

    # path joining version for other paths
    image_count = len([name for name in os.listdir(image_path) if os.path.isfile(os.path.join(image_path, name))])

    return document_text, word_count, paragraph_count, bullet_count, image_count


def analyseDocxFilePypandoc(fileName):
    # https://stackoverflow.com/questions/62859658/how-to-convert-docx-to-txt-in-python
    # https://pypi.org/project/pypandoc/
    # https://pandoc.org/installing.html
    word_count = 0
    paragraph_count = 0
    bullet_count = 0
    image_count = 0

    # create temp text file and populate with document text
    document = fileName
    temp_file = './data/word_text.txt'
    converted_document = pypandoc.convert_file(document, 'plain', outputfile=temp_file)
    assert converted_document == ""

    # read text from temp text file for analysis
    document = open(temp_file, encoding='utf8')
    document_text = document.read()
    document_words = document_text.split()
    word_count = len(document_words)

    paragraph_count = document_text.count('\n\n')

    bullet_count = document_text.count('-   ')
    
    return document_text, word_count, paragraph_count, bullet_count, image_count

